"""
13주차 docx detector + guide builder 단위 테스트

테스트 케이스 TC1~TC18을 실행합니다.

테스트는 다음 두 함수를 검증합니다.
- detect_in_docx(): docx 파일을 탐지하여 DeidentifyPlan 생성
- build_guide_for_docx(): DeidentifyPlan을 guide 모드 CommonApplyResult로 변환

탐지 함수는 mock으로 주입해 모델 의존성을 끊었습니다.
"""

from __future__ import annotations

from pathlib import Path
import sys

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = PROJECT_ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from docx import Document

from common_apply_result import APPLY_MODE_GUIDE
from deidentify_target_builder import DeidentifyPlan, DeidentifyTarget
from docx_detector import (
    build_guide_for_docx,
    detect_and_build_guide_for_docx,
    detect_in_docx,
)


# ── 테스트용 docx 생성 헬퍼 ────────────────────────────────────

def _create_sample_docx(paragraphs: list[str], path: Path) -> Path:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _create_docx_with_table(
    body_paragraphs: list[str],
    table_data: list[list[str]],
    path: Path,
) -> Path:
    """본문 + 표가 들어있는 docx 생성.

    table_data: [[행1셀1, 행1셀2], [행2셀1, 행2셀2], ...]
    """
    doc = Document()
    for text in body_paragraphs:
        doc.add_paragraph(text)
    if table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        table = doc.add_table(rows=rows, cols=cols)
        for r, row_data in enumerate(table_data):
            for c, cell_text in enumerate(row_data):
                table.cell(r, c).text = cell_text
    doc.save(str(path))
    return path


def _create_docx_with_merged_cell(path: Path) -> Path:
    """병합 셀이 있는 docx 생성 (1행에 가로 병합).

    구조:
      [병합된 셀 (1행 1-2열)]
      [셀(2,1)] [셀(2,2)]
    """
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # 1행의 두 셀을 가로 병합
    a = table.cell(0, 0)
    b = table.cell(0, 1)
    a.merge(b)
    a.text = "병합 셀 이메일: merged@example.com"
    table.cell(1, 0).text = "셀(2,1) 일반"
    table.cell(1, 1).text = "셀(2,2) 일반"
    doc.save(str(path))
    return path


def _make_target(
    *,
    label: str,
    matched: str,
    start: int | None,
    end: int | None,
    source: str,
    action: str,
    paragraph_no: int,
    section: str = "body",
    context: str,
    file_type: str = "docx",
    grade: str = "S",
    order: int = 0,
    table_no: int | None = None,
    row_no: int | None = None,
    col_no: int | None = None,
) -> DeidentifyTarget:
    location_meta: dict = {
        "fileType": file_type,
        "section": section,
        "paragraphNo": paragraph_no,
    }
    if table_no is not None:
        location_meta["tableNo"] = table_no
    if row_no is not None:
        location_meta["rowNo"] = row_no
    if col_no is not None:
        location_meta["colNo"] = col_no

    return DeidentifyTarget(
        label=label,
        matched=matched,
        action=action,
        location_label=None,  # locationLabel 자동 생성 확인용
        location_meta=location_meta,
        start=start,
        end=end,
        source=source,
        reason=f"테스트용 {source} 탐지",
        grade=grade,
        context=context,
        order=order,
    )


# ── 테스트 헬퍼 ─────────────────────────────────────────────────

sys.path.insert(0, str(Path(__file__).resolve().parent))
from test_helpers import TestRunner, run_test_functions

_runner = TestRunner("13주차 docx detector 단위 테스트")


def _check(tc_id: str, condition: bool, message: str = "") -> None:
    _runner.check(tc_id, condition, message)


# ── TC1: 이메일 paragraph 1개 ──────────────────────────────────

def tc1(tmp_dir: Path) -> None:
    print("\nTC1: 이메일이 있는 paragraph 1개 마스킹")

    text = "담당자 이메일은 test@example.com입니다."
    path = _create_sample_docx([text], tmp_dir / "tc1.docx")

    # start: "test@example.com" 위치
    matched = "test@example.com"
    start = text.index(matched)
    end = start + len(matched)

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(
                label="이메일 주소", matched=matched, start=start, end=end,
                source="regex", action="마스킹", paragraph_no=0, context=text,
            )
        ],
        review_targets=[],
    )

    result = build_guide_for_docx(str(path), plan)

    _check("TC1.applyMode", result.applyMode == APPLY_MODE_GUIDE)
    _check("TC1.outputFilePath", result.outputFilePath is None)
    _check("TC1.autoResults_count", len(result.autoResults) == 1)
    item = result.autoResults[0]
    _check("TC1.status_applied", item.status == "applied")
    _check("TC1.applied_count", item.appliedTargetCount == 1)
    _check("TC1.appliedText", "*" * len(matched) in item.appliedText)


# ── TC2: 성명 paragraph 1개 ─────────────────────────────────────

def tc2(tmp_dir: Path) -> None:
    print("\nTC2: 성명이 있는 paragraph 1개 마스킹")

    text = "직원 김도윤의 서류를 검토했습니다."
    path = _create_sample_docx([text], tmp_dir / "tc2.docx")

    matched = "김도윤"
    start = text.index(matched)
    end = start + len(matched)

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(
                label="성명", matched=matched, start=start, end=end,
                source="ner", action="마스킹", paragraph_no=0, context=text,
            )
        ],
        review_targets=[],
    )

    result = build_guide_for_docx(str(path), plan)
    item = result.autoResults[0]
    _check("TC2.status", item.status == "applied")
    _check("TC2.appliedText", "***" in item.appliedText)


# ── TC3: 성명 + 이메일 동시 ─────────────────────────────────────

def tc3(tmp_dir: Path) -> None:
    print("\nTC3: 한 paragraph에 성명 + 이메일 동시 마스킹")

    text = "담당자 김도윤의 이메일은 test@example.com입니다."
    path = _create_sample_docx([text], tmp_dir / "tc3.docx")

    s1 = text.index("김도윤"); e1 = s1 + 3
    s2 = text.index("test@example.com"); e2 = s2 + len("test@example.com")

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(label="성명", matched="김도윤", start=s1, end=e1,
                         source="ner", action="마스킹", paragraph_no=0, context=text, order=0),
            _make_target(label="이메일 주소", matched="test@example.com", start=s2, end=e2,
                         source="regex", action="마스킹", paragraph_no=0, context=text, order=1),
        ],
        review_targets=[],
    )

    result = build_guide_for_docx(str(path), plan)
    _check("TC3.autoResults_count", len(result.autoResults) == 1)
    item = result.autoResults[0]
    _check("TC3.applied_count", item.appliedTargetCount == 2)
    _check("TC3.status", item.status == "applied")
    _check("TC3.label_order",
           item.label == "성명, 이메일 주소",
           f"label={item.label}")


# ── TC4: 내부 IP 삭제 권장 ──────────────────────────────────────

def tc4(tmp_dir: Path) -> None:
    print("\nTC4: 내부 IP 삭제 권장")

    text = "서버 IP는 192.168.0.1입니다."
    path = _create_sample_docx([text], tmp_dir / "tc4.docx")

    matched = "192.168.0.1"
    start = text.index(matched); end = start + len(matched)

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(label="내부 IP 주소", matched=matched, start=start, end=end,
                         source="regex", action="삭제", paragraph_no=0, context=text, grade="C"),
        ],
        review_targets=[],
    )

    # default: deletion_mode="delete"
    result = build_guide_for_docx(str(path), plan)
    item = result.autoResults[0]
    _check("TC4.delete_mode", matched not in item.appliedText)
    _check("TC4.applied_text", item.appliedText == "서버 IP는 입니다.")

    # mark mode
    result_mark = build_guide_for_docx(str(path), plan, deletion_mode="mark")
    item_mark = result_mark.autoResults[0]
    _check("TC4.mark_mode", "(삭제됨)" in item_mark.appliedText)


# ── TC5: reviewTargets 보존 ─────────────────────────────────────

def tc5(tmp_dir: Path) -> None:
    print("\nTC5: reviewTargets 보존")

    text = "입찰 제안 평가표를 검토했습니다."
    path = _create_sample_docx([text], tmp_dir / "tc5.docx")

    review = _make_target(
        label="민감정보", matched="", start=None, end=None,
        source="ai", action="검토 필요", paragraph_no=0, context=text, grade="C",
    )

    plan = DeidentifyPlan(auto_targets=[], review_targets=[review])

    result = build_guide_for_docx(str(path), plan)
    _check("TC5.review_count", len(result.reviewTargets) == 1)
    _check("TC5.autoResults_empty", len(result.autoResults) == 0)
    rv = result.reviewTargets[0]
    _check("TC5.review_label", rv.label == "민감정보")
    _check("TC5.review_action", rv.action == "검토 필요")


# ── TC6: paragraphNo 없음 ───────────────────────────────────────

def tc6(tmp_dir: Path) -> None:
    print("\nTC6: paragraphNo 없음")

    text = "담당자 이메일은 test@example.com입니다."
    path = _create_sample_docx([text], tmp_dir / "tc6.docx")

    target = DeidentifyTarget(
        label="이메일 주소", matched="test@example.com",
        action="마스킹", location_label="알 수 없음",
        location_meta={"fileType": "docx", "section": "body"},  # paragraphNo 누락
        start=0, end=16, source="regex", reason="test", grade="S", context=text,
    )

    plan = DeidentifyPlan(auto_targets=[target], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    _check("TC6.skipped_count", len(result.autoResults) == 1)
    item = result.autoResults[0]
    _check("TC6.status", item.status == "skipped")
    _check("TC6.warning_type",
           any("[missing_paragraph_no]" in w for w in item.warnings),
           f"warnings={item.warnings}")


# ── TC7: paragraphNo 범위 초과 ──────────────────────────────────

def tc7(tmp_dir: Path) -> None:
    print("\nTC7: paragraphNo 범위 초과")

    text = "단 하나의 문단."
    path = _create_sample_docx([text], tmp_dir / "tc7.docx")

    target = _make_target(
        label="성명", matched="홍길동", start=0, end=3,
        source="ner", action="마스킹", paragraph_no=99,  # 범위 초과
        context="홍길동 무관 텍스트",
    )

    plan = DeidentifyPlan(auto_targets=[target], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    item = result.autoResults[0]
    _check("TC7.status", item.status == "skipped")
    _check("TC7.warning_type",
           any("[paragraph_out_of_range]" in w for w in item.warnings))


# ── TC8: context 불일치, slice 일치 ─────────────────────────────

def tc8(tmp_dir: Path) -> None:
    print("\nTC8: context 불일치, slice 일치 → 권장 + warning")

    text = "담당자 이메일은 test@example.com입니다."
    path = _create_sample_docx([text], tmp_dir / "tc8.docx")

    matched = "test@example.com"
    start = text.index(matched); end = start + len(matched)

    # context를 의도적으로 다르게
    target = _make_target(
        label="이메일 주소", matched=matched, start=start, end=end,
        source="regex", action="마스킹", paragraph_no=0,
        context="담당자 이메일: test@example.com",
    )

    plan = DeidentifyPlan(auto_targets=[target], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    item = result.autoResults[0]
    _check("TC8.status_applied", item.status == "applied")
    _check("TC8.applied_count", item.appliedTargetCount == 1)
    _check("TC8.context_mismatch_warning",
           any("[context_mismatch]" in w for w in item.warnings))


# ── TC9: slice 불일치 ──────────────────────────────────────────

def tc9(tmp_dir: Path) -> None:
    print("\nTC9: slice 불일치 → skip")

    text = "담당자 이메일은 test@example.com입니다."
    path = _create_sample_docx([text], tmp_dir / "tc9.docx")

    # matched와 다른 위치 지정
    target = _make_target(
        label="이메일 주소", matched="test@example.com",
        start=0, end=5,  # 일치하지 않는 위치
        source="regex", action="마스킹", paragraph_no=0, context=text,
    )

    plan = DeidentifyPlan(auto_targets=[target], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    item = result.autoResults[0]
    _check("TC9.status_skipped", item.status == "skipped")
    _check("TC9.skipped_count", item.skippedTargetCount == 1)
    _check("TC9.slice_mismatch_warning",
           any("[slice_mismatch]" in w for w in item.warnings))


# ── TC10: 여러 paragraph에 target 분산 ──────────────────────────

def tc10(tmp_dir: Path) -> None:
    print("\nTC10: 여러 paragraph에 target 분산")

    texts = [
        "첫 번째 문단입니다.",
        "담당자 김도윤의 서류.",
        "이메일은 test@example.com입니다.",
    ]
    path = _create_sample_docx(texts, tmp_dir / "tc10.docx")

    s1 = texts[1].index("김도윤"); e1 = s1 + 3
    s2 = texts[2].index("test@example.com"); e2 = s2 + len("test@example.com")

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(label="성명", matched="김도윤", start=s1, end=e1,
                         source="ner", action="마스킹", paragraph_no=1, context=texts[1]),
            _make_target(label="이메일 주소", matched="test@example.com", start=s2, end=e2,
                         source="regex", action="마스킹", paragraph_no=2, context=texts[2]),
        ],
        review_targets=[],
    )

    result = build_guide_for_docx(str(path), plan)
    _check("TC10.autoResults_count", len(result.autoResults) == 2)
    statuses = [item.status for item in result.autoResults]
    _check("TC10.all_applied", all(s == "applied" for s in statuses),
           f"statuses={statuses}")


# ── TC11: summary 정합성 ────────────────────────────────────────

def tc11(tmp_dir: Path) -> None:
    print("\nTC11: summary 정합성")

    texts = ["담당자 김도윤입니다.", "이메일: test@example.com"]
    path = _create_sample_docx(texts, tmp_dir / "tc11.docx")

    s1 = texts[0].index("김도윤"); e1 = s1 + 3
    s2 = texts[1].index("test@example.com"); e2 = s2 + len("test@example.com")

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(label="성명", matched="김도윤", start=s1, end=e1,
                         source="ner", action="마스킹", paragraph_no=0, context=texts[0]),
            _make_target(label="이메일 주소", matched="test@example.com", start=s2, end=e2,
                         source="regex", action="마스킹", paragraph_no=1, context=texts[1]),
        ],
        review_targets=[],
    )

    result = build_guide_for_docx(str(path), plan)
    s = result.summary
    _check("TC11.total", s.totalLocations == len(result.autoResults))
    _check("TC11.sum_breakdown",
           s.appliedLocations + s.partialLocations + s.skippedLocations == s.totalLocations)
    sum_targets = sum(
        item.appliedTargetCount + item.skippedTargetCount
        for item in result.autoResults
    )
    _check("TC11.auto_target_count", s.autoTargetCount == sum_targets,
           f"summary={s.autoTargetCount}, computed={sum_targets}")


# ── TC12: deletion_mode=mark ────────────────────────────────────

def tc12(tmp_dir: Path) -> None:
    print("\nTC12: deletion_mode=mark")

    text = "VLAN 100을 사용합니다."
    path = _create_sample_docx([text], tmp_dir / "tc12.docx")

    matched = "VLAN 100"
    start = text.index(matched); end = start + len(matched)

    plan = DeidentifyPlan(
        auto_targets=[
            _make_target(label="VLAN", matched=matched, start=start, end=end,
                         source="regex", action="삭제", paragraph_no=0, context=text, grade="C"),
        ],
        review_targets=[],
    )

    result = build_guide_for_docx(str(path), plan, deletion_mode="mark")
    item = result.autoResults[0]
    _check("TC12.mark", "(삭제됨)" in item.appliedText)


# ── TC13: 빈 paragraph는 detection 대상 제외 ───────────────────

def tc13(tmp_dir: Path) -> None:
    print("\nTC13: 빈 paragraph는 detection 대상 제외")

    # docx 본문에 빈 문단을 섞어 둠
    doc = Document()
    doc.add_paragraph("")  # 빈 문단
    doc.add_paragraph("담당자 이메일은 test@example.com입니다.")
    doc.add_paragraph("")  # 빈 문단
    path = tmp_dir / "tc13.docx"
    doc.save(str(path))

    # mock regex detector: 이메일 한 개 반환
    def mock_regex(text: str):
        if "test@example.com" in text:
            idx = text.index("test@example.com")
            return [{
                "label": "이메일 주소",
                "value": "test@example.com",
                "start": idx, "end": idx + len("test@example.com"),
                "grade": "S", "action": "마스킹", "desc": "test",
            }]
        return []

    plan = detect_in_docx(str(path), regex_detect_func=mock_regex)
    _check("TC13.detection_count", len(plan.auto_targets) == 1)
    # paragraphNo는 1이어야 함 (빈 문단 포함 인덱스)
    _check("TC13.paragraphNo",
           plan.auto_targets[0].location_meta.get("paragraphNo") == 1,
           f"paragraphNo={plan.auto_targets[0].location_meta.get('paragraphNo')}")


# ── TC14: locationLabel 형식 ────────────────────────────────────

def tc14(tmp_dir: Path) -> None:
    print('\nTC14: locationLabel에 "본문 N번째 문단:" 형식 + context')

    long_text = "담당자 김도윤의 이메일은 test@example.com입니다. 길게 이어지는 문장."
    path = _create_sample_docx([long_text], tmp_dir / "tc14.docx")

    def mock_regex(text: str):
        return []  # 빈 결과여도 paragraph는 순회됨

    plan = detect_in_docx(str(path), regex_detect_func=mock_regex)
    # detection은 없지만, paragraph 라벨 생성 자체를 별도 확인
    from docx_detector import iter_body_paragraphs, load_docx
    paragraphs = iter_body_paragraphs(load_docx(str(path)))
    label = paragraphs[0].location_label
    _check("TC14.prefix", label.startswith("본문 1번째 문단:"), f"label={label}")
    _check("TC14.has_context", "담당자" in label)
    _check("TC14.length_limit", "..." in label, f"label={label}")


# ── TC15: applyMode 및 outputFilePath ──────────────────────────

def tc15(tmp_dir: Path) -> None:
    print("\nTC15: applyMode=\"guide\", outputFilePath=None")

    text = "단순 문단입니다."
    path = _create_sample_docx([text], tmp_dir / "tc15.docx")

    plan = DeidentifyPlan(auto_targets=[], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    _check("TC15.applyMode", result.applyMode == "guide")
    _check("TC15.outputFilePath_None", result.outputFilePath is None)
    _check("TC15.fileType", result.fileType == "docx")


# ── TC16: 공백/탭만 있는 paragraph ──────────────────────────────

def tc16(tmp_dir: Path) -> None:
    print("\nTC16: 공백/탭만 있는 paragraph는 TextUnit 생성 안 됨")

    doc = Document()
    doc.add_paragraph("   ")        # 공백만
    doc.add_paragraph("\t\t")       # 탭만
    doc.add_paragraph("실제 텍스트.")
    path = tmp_dir / "tc16.docx"
    doc.save(str(path))

    from docx_detector import iter_body_paragraphs, load_docx
    paragraphs = iter_body_paragraphs(load_docx(str(path)))
    _check("TC16.paragraph_count", len(paragraphs) == 1)
    _check("TC16.paragraphNo", paragraphs[0].paragraph_no == 2,
           f"paragraphNo={paragraphs[0].paragraph_no}")


# ── TC17: section이 body가 아닌 target ──────────────────────────

def tc17(tmp_dir: Path) -> None:
    print("\nTC17: section이 body가 아닌 target → skip")

    text = "본문 문단입니다."
    path = _create_sample_docx([text], tmp_dir / "tc17.docx")

    target = _make_target(
        label="성명", matched="홍길동", start=0, end=3,
        source="ner", action="마스킹", paragraph_no=0, section="table_cell",
        context="홍길동 데이터",
    )

    plan = DeidentifyPlan(auto_targets=[target], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    item = result.autoResults[0]
    _check("TC17.status", item.status == "skipped")
    _check("TC17.warning_type",
           any("[missing_table_cell_location]" in w for w in item.warnings))


# ── TC18: target 위치 겹침 ──────────────────────────────────────

def tc18(tmp_dir: Path) -> None:
    print("\nTC18: 동일 paragraph 내 target 위치가 겹침")

    text = "담당자 김도윤의 서류."
    path = _create_sample_docx([text], tmp_dir / "tc18.docx")

    s = text.index("김도윤"); e = s + 3

    # 같은 구간을 regex와 ner가 동시에 잡은 경우
    # DeidentifyPlan 생성 단계에서 source 우선순위(regex>ner)로 처리되어야 함
    detections = [
        {
            "label": "이름 패턴", "matched": "김도윤",
            "grade": "S", "action": "마스킹", "source": "regex",
            "context": text, "locationLabel": None,
            "locationMeta": {"fileType": "docx", "section": "body", "paragraphNo": 0},
            "start": s, "end": e, "_order": 0,
        },
        {
            "label": "성명", "matched": "김도윤",
            "grade": "S", "action": "마스킹", "source": "ner",
            "context": text, "locationLabel": None,
            "locationMeta": {"fileType": "docx", "section": "body", "paragraphNo": 0},
            "start": s, "end": e, "_order": 1,
        },
    ]

    from deidentify_target_builder import build_deidentify_plan
    plan = build_deidentify_plan(detections)

    _check("TC18.dedup_to_one", len(plan.auto_targets) == 1,
           f"count={len(plan.auto_targets)}")
    _check("TC18.regex_kept", plan.auto_targets[0].source == "regex",
           f"source={plan.auto_targets[0].source}")

    # guide 생성도 정상 동작
    result = build_guide_for_docx(str(path), plan)
    item = result.autoResults[0]
    _check("TC18.guide_applied", item.status == "applied")
    _check("TC18.applied_count", item.appliedTargetCount == 1)


# ── TC19: docx 표 셀 단일 탐지 ────────────────────────────────

def tc19(tmp_dir: Path) -> None:
    print("\nTC19: 표 셀 안의 이메일 탐지 + guide")

    body = ["본문 paragraph"]
    table_data = [
        ["담당자", "이메일: test@example.com"],
        ["연락처", "010-1234-5678"],
    ]
    path = _create_docx_with_table(body, table_data, tmp_dir / "tc19.docx")

    matched = "test@example.com"
    cell_text = "이메일: test@example.com"
    s = cell_text.index(matched); e = s + len(matched)

    target = _make_target(
        label="이메일 주소", matched=matched, start=s, end=e,
        source="regex", action="마스킹",
        paragraph_no=0, section="table_cell",
        table_no=0, row_no=0, col_no=1,
        context=cell_text,
    )

    plan = DeidentifyPlan(auto_targets=[target], review_targets=[])
    result = build_guide_for_docx(str(path), plan)

    _check("TC19.autoResults_count", len(result.autoResults) == 1)
    item = result.autoResults[0]
    _check("TC19.status", item.status == "applied")
    _check("TC19.applied_count", item.appliedTargetCount == 1)
    _check("TC19.preview_masked", "*" * len(matched) in item.appliedText)
    _check("TC19.meta_section",
           item.locationMeta.get("section") == "table_cell")
    _check("TC19.meta_table_no",
           item.locationMeta.get("tableNo") == 0)
    _check("TC19.meta_row_col",
           item.locationMeta.get("rowNo") == 0
           and item.locationMeta.get("colNo") == 1)


# ── TC20: 본문 + 표 동시 탐지 ─────────────────────────────────

def tc20(tmp_dir: Path) -> None:
    print("\nTC20: 본문 + 표 셀 동시 탐지")

    body = ["본문: 홍길동 직원 정보"]
    table_data = [
        ["사번", "이메일: foo@example.com"],
    ]
    path = _create_docx_with_table(body, table_data, tmp_dir / "tc20.docx")

    # 본문 target (paragraph 0)
    body_text = body[0]
    name_idx = body_text.index("홍길동")
    body_target = _make_target(
        label="성명", matched="홍길동", start=name_idx, end=name_idx + 3,
        source="ner", action="마스킹",
        paragraph_no=0, section="body",
        context=body_text, order=0,
    )

    # 표 target
    cell_text = "이메일: foo@example.com"
    email_idx = cell_text.index("foo@example.com")
    table_target = _make_target(
        label="이메일", matched="foo@example.com",
        start=email_idx, end=email_idx + 15,
        source="regex", action="마스킹",
        paragraph_no=0, section="table_cell",
        table_no=0, row_no=0, col_no=1,
        context=cell_text, order=1,
    )

    plan = DeidentifyPlan(
        auto_targets=[body_target, table_target],
        review_targets=[],
    )
    result = build_guide_for_docx(str(path), plan)

    _check("TC20.two_items", len(result.autoResults) == 2)
    sections = sorted(
        item.locationMeta.get("section") for item in result.autoResults
    )
    _check("TC20.sections_mixed", sections == ["body", "table_cell"])

    # 둘 다 applied
    statuses = {item.status for item in result.autoResults}
    _check("TC20.all_applied", statuses == {"applied"})


# ── TC21: 다중 표/다중 셀 탐지 ────────────────────────────────

def tc21(tmp_dir: Path) -> None:
    print("\nTC21: 표 2개, 각 표의 여러 셀에서 탐지")

    body: list[str] = []

    doc_path = tmp_dir / "tc21.docx"
    doc = Document()
    # 표 1: 2x2
    t1 = doc.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "이름"
    t1.cell(0, 1).text = "홍길동"
    t1.cell(1, 0).text = "사번"
    t1.cell(1, 1).text = "12345"
    # 표 2: 1x2
    t2 = doc.add_table(rows=1, cols=2)
    t2.cell(0, 0).text = "이메일"
    t2.cell(0, 1).text = "test@example.com"
    doc.save(str(doc_path))

    # 3건 탐지: 표1(0,1), 표1(1,1), 표2(0,1)
    targets = [
        _make_target(
            label="성명", matched="홍길동", start=0, end=3,
            source="ner", action="마스킹",
            paragraph_no=0, section="table_cell",
            table_no=0, row_no=0, col_no=1,
            context="홍길동", order=0,
        ),
        _make_target(
            label="사번", matched="12345", start=0, end=5,
            source="regex", action="마스킹",
            paragraph_no=0, section="table_cell",
            table_no=0, row_no=1, col_no=1,
            context="12345", order=1,
        ),
        _make_target(
            label="이메일", matched="test@example.com", start=0, end=16,
            source="regex", action="마스킹",
            paragraph_no=0, section="table_cell",
            table_no=1, row_no=0, col_no=1,
            context="test@example.com", order=2,
        ),
    ]

    plan = DeidentifyPlan(auto_targets=targets, review_targets=[])
    result = build_guide_for_docx(str(doc_path), plan)

    _check("TC21.three_items", len(result.autoResults) == 3)

    table_nos = sorted({
        item.locationMeta.get("tableNo") for item in result.autoResults
    })
    _check("TC21.two_tables", table_nos == [0, 1])

    statuses = {item.status for item in result.autoResults}
    _check("TC21.all_applied", statuses == {"applied"})


# ── TC22: 병합 셀이 있어도 일반 셀 누락 없음 ──────────────────

def tc22(tmp_dir: Path) -> None:
    print("\nTC22: 병합 셀이 있어도 일반 셀이 누락되지 않음 (13주차 정책)")

    # 1행 가로 병합, 2행은 일반 두 셀
    # 13주차 정책: seen_cells 제거. 중복 안내보다 탐지 누락 방지 우선.
    path = _create_docx_with_merged_cell(tmp_dir / "tc22.docx")

    def mock_regex(text):
        import re
        results = []
        # 일반 셀에서만 탐지될 패턴: "셀(N,M)"
        for m in re.finditer(r'셀\(\d,\d\)', text):
            results.append({
                "label": "셀_식별자", "value": m.group(),
                "start": m.start(), "end": m.end(),
                "grade": "S", "action": "마스킹", "desc": "셀",
            })
        return results

    plan = detect_in_docx(str(path), regex_detect_func=mock_regex)

    # 2행의 일반 셀 두 개에서 모두 탐지되어야 함
    cell_targets = [t for t in plan.auto_targets if t.label == "셀_식별자"]
    cell_texts = sorted({t.matched for t in cell_targets})

    _check("TC22.both_normal_cells_detected",
           cell_texts == ["셀(2,1)", "셀(2,2)"],
           f"탐지된 셀 텍스트: {cell_texts}")

    # 두 일반 셀이 각각 (1,0)과 (1,1)에 위치
    row1_cols = sorted({
        t.location_meta.get("colNo") for t in cell_targets
        if t.location_meta.get("rowNo") == 1
    })
    _check("TC22.row_1_both_cols",
           row1_cols == [0, 1],
           f"row=1 col={row1_cols}")


# ── 실행 ──────────────────────────────────────────────────────

def main() -> None:
    import tempfile
    with tempfile.TemporaryDirectory() as tmp_dir_str:
        tmp_dir = Path(tmp_dir_str)
        print("=== 13주차 docx detector 단위 테스트 ===")
        test_fns = [
            tc1, tc2, tc3, tc4, tc5, tc6, tc7, tc8, tc9,
            tc10, tc11, tc12, tc13, tc14, tc15, tc16, tc17, tc18,
            tc19, tc20, tc21, tc22,
        ]
        run_test_functions(_runner, test_fns, tmp_dir)

    _runner.report()


if __name__ == "__main__":
    main()
