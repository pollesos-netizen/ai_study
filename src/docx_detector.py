"""
docx ?뚯씪 ?먯? + ?덈궡(guide) 紐⑤뱶

紐⑹쟻:
- docx ?뚯씪?먯꽌 蹂몃Ц paragraph瑜??쒗쉶?섎ŉ 媛쒖씤?뺣낫/誘쇨컧?뺣낫瑜??먯??⑸땲??
- ?쒖뒪?쒖? docx ?뚯씪??吏곸젒 ?섏젙?섏? ?딆뒿?덈떎.
- 寃곌낵??CommonApplyResult(applyMode="guide")濡?諛섑솚?⑸땲??
- ?ъ슜?먮뒗 ?덈궡???곕씪 ?먮낯 docx?먯꽌 吏곸젒 ?섏젙?⑸땲??

13二쇱감 踰붿쐞:
- 蹂몃Ц paragraphs留?泥섎━ (???ㅻ뜑/?명꽣/媛곸＜??蹂꾨룄 二쇱감)
- 鍮?臾몃떒(strip 湲곗?)? ?먯? ??곸뿉???쒖쇅
- regex + NER + AI ?먯?瑜?吏?먰븯?? 13二쇱감 珥덈컲?먮뒗 regex留??곗꽑 寃利?媛??
?듭떖 ?⑥닔:
- detect_in_docx():            DeidentifyPlan ?앹꽦
- build_guide_for_docx():      DeidentifyPlan -> CommonApplyResult (guide)
- detect_and_build_guide_for_docx(): ???⑥닔???몄쓽 wrapper
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

try:
    from src.common_apply_result import (
        APPLY_MODE_GUIDE,
        CommonApplyItem,
        CommonApplyResult,
        build_summary,
        make_review_items,
    )
    from src.common_apply_utils import (
        WARNING_CONTEXT_MISMATCH,
        WARNING_EMPTY_PARAGRAPH_TARGET,
        WARNING_MISSING_PARAGRAPH_NO,
        WARNING_UNSUPPORTED_DOCX_SECTION,
        WARNING_PARAGRAPH_OUT_OF_RANGE,
        actions_for_targets,
        format_warning,
        labels_for_targets,
        make_location_label_with_context,
        make_status,
        normalize_nfc,
        validate_slice_against_text,
    )
    from src.deidentify_apply import apply_targets_to_text
    from src.deidentify_target_builder import (
        DeidentifyPlan,
        DeidentifyTarget,
        build_deidentify_plan,
    )
except ModuleNotFoundError:
    from common_apply_result import (
        APPLY_MODE_GUIDE,
        CommonApplyItem,
        CommonApplyResult,
        build_summary,
        make_review_items,
    )
    from common_apply_utils import (
        WARNING_CONTEXT_MISMATCH,
        WARNING_EMPTY_PARAGRAPH_TARGET,
        WARNING_MISSING_PARAGRAPH_NO,
        WARNING_UNSUPPORTED_DOCX_SECTION,
        WARNING_PARAGRAPH_OUT_OF_RANGE,
        actions_for_targets,
        format_warning,
        labels_for_targets,
        make_location_label_with_context,
        make_status,
        normalize_nfc,
        validate_slice_against_text,
    )
    from deidentify_apply import apply_targets_to_text
    from deidentify_target_builder import (
        DeidentifyPlan,
        DeidentifyTarget,
        build_deidentify_plan,
    )


# ?? ?곗씠??援ъ“ ????????????????????????????????????????????????

@dataclass
class ParsedParagraph:
    """
    docx??paragraph瑜??곕━ ?먯? ?⑥쐞濡?蹂?섑븳 援ъ“.

    section:
    - "body": doc.paragraphs 湲곗? 蹂몃Ц paragraph
    - "table_cell": doc.tables ?대? cell.paragraphs 湲곗? paragraph
    """

    paragraph_no: int  # section蹂?paragraph ?몃뜳??(0-based)
    section: str
    text: str
    table_no: int | None = None
    row_no: int | None = None
    col_no: int | None = None

    @property
    def location_label(self) -> str:
        # ?ъ슜???쒖떆留?1-based濡?蹂??+ context 30??        if self.section == "table_cell":
            base = (
                f"??{self.table_no + 1}踰?"
                f"{self.row_no + 1}??{self.col_no + 1}??
            )
            if self.paragraph_no > 0:
                base += f" {self.paragraph_no + 1}踰덉㎏ 臾몃떒"
        else:
            base = f"蹂몃Ц {self.paragraph_no + 1}踰덉㎏ 臾몃떒"

        return make_location_label_with_context(base, self.text, max_length=30)

    @property
    def location_meta(self) -> dict[str, Any]:
        meta = {
            "fileType": "docx",
            "section": self.section,
            "paragraphNo": self.paragraph_no,
        }

        if self.section == "table_cell":
            meta.update({
                "tableNo": self.table_no,
                "rowNo": self.row_no,
                "colNo": self.col_no,
            })

        return meta


# ?? docx 濡쒕뱶 諛?paragraph ?쒗쉶 ????????????????????????????????

def load_docx(input_path: str | Path):
    """
    python-docx濡?docx ?뚯씪??濡쒕뱶?⑸땲??

    ?섏〈?? python-docx (`pip install python-docx`)
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx媛 ?ㅼ튂?섏뼱 ?덉? ?딆뒿?덈떎. "
            "`pip install python-docx`瑜??ㅽ뻾?섏꽭??"
        ) from exc

    return Document(str(input_path))


def iter_body_paragraphs(doc) -> list[ParsedParagraph]:
    """
    臾몄꽌 蹂몃Ц(doc.paragraphs)??paragraph瑜?ParsedParagraph 紐⑸줉?쇰줈 諛섑솚?⑸땲??

    - 鍮?臾몃떒(strip() 湲곗?)? ?쒖쇅?⑸땲??
    - paragraphNo??鍮?臾몃떒???ы븿???먮Ц ?몃뜳?ㅻ? ?좎??⑸땲??
    - 13二쇱감?먯꽌??蹂몃Ц留?泥섎━?섎?濡?section="body"濡?怨좎젙?⑸땲??
    """
    parsed: list[ParsedParagraph] = []

    for para_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text

        if not text.strip():
            continue

        parsed.append(
            ParsedParagraph(
                paragraph_no=para_index,
                section="body",
                text=text,
            )
        )

    return parsed


def iter_table_cell_paragraphs(doc) -> list[ParsedParagraph]:
    """
    臾몄꽌 ??? ?대???paragraph瑜?ParsedParagraph 紐⑸줉?쇰줈 諛섑솚?⑸땲??

    - 鍮?paragraph(strip() 湲곗?)???쒖쇅?⑸땲??
    - paragraphNo???대떦 cell.paragraphs 湲곗? ?몃뜳?ㅻ? ?좎??⑸땲??
    - 蹂묓빀 ?? python-docx?먯꽌 媛숈? XML cell??以묐났 李몄“?????덉쑝誘濡?      id(cell._tc) 湲곗??쇰줈 以묐났 ?먯?瑜?諛⑹??⑸땲??
    """
    parsed: list[ParsedParagraph] = []
    seen_cells: set[int] = set()

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)

                for para_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text
                    if not text.strip():
                        continue

                    parsed.append(
                        ParsedParagraph(
                            paragraph_no=para_index,
                            section="table_cell",
                            text=text,
                            table_no=table_index,
                            row_no=row_index,
                            col_no=col_index,
                        )
                    )

    return parsed


def iter_docx_paragraphs(doc) -> list[ParsedParagraph]:
    """
    13二쇱감 docx guide ?먯? ???paragraph瑜?諛섑솚?⑸땲??

    ???
    - 蹂몃Ц paragraph
    - ??? ?대? paragraph

    ?쒖쇅:
    - ?ㅻ뜑/?명꽣/媛곸＜/二쇱꽍/?꾪삎/SmartArt/李⑦듃 ?대? ?띿뒪??    """
    return iter_body_paragraphs(doc) + iter_table_cell_paragraphs(doc)


# ?? Detection ?앹꽦 (regex / NER / AI ?대뙌?? ???????????????????

def _make_target_dict_from_regex(
    raw: Any,
    paragraph: ParsedParagraph,
    order: int,
) -> dict[str, Any] | None:
    """
    regex_detector??寃곌낵(DetectionResult ?먮뒗 ?좎궗 dict)瑜?Detection dict濡?蹂?섑빀?덈떎.

    DeidentifyPlan ?앹꽦湲곕뒗 dict ?낅젰??諛쏆쑝誘濡?dict ?뺥깭濡??뺢퇋?뷀빀?덈떎.
    """
    def _get(obj, *names, default=None):
        for name in names:
            if isinstance(obj, dict) and name in obj:
                return obj.get(name)
            if hasattr(obj, name):
                return getattr(obj, name)
        return default

    label = _get(raw, "label")
    value = _get(raw, "value", "matched", "match")
    start = _get(raw, "start")
    end = _get(raw, "end")
    grade = _get(raw, "grade", default="S")
    action = _get(raw, "action", default="留덉뒪??)
    desc = _get(raw, "desc", "reason", default=None)

    if label is None or value is None or start is None or end is None:
        return None

    return {
        "label": str(label),
        "matched": str(value),
        "grade": str(grade),
        "action": str(action),
        "source": "regex",
        "context": paragraph.text,
        "locationLabel": paragraph.location_label,
        "locationMeta": paragraph.location_meta,
        "start": int(start),
        "end": int(end),
        "sensitiveType": _get(raw, "sensitive_type", "sensitiveType", default=None),
        "sensitiveCategory": _get(raw, "sensitive_category", "sensitiveCategory", default=label),
        "reason": str(desc) if desc else f"?뺢퇋???먯?: {label}",
        "_order": order,
    }


def _make_target_dict_from_ner(
    raw: dict[str, Any],
    paragraph: ParsedParagraph,
    order: int,
    *,
    threshold: float,
) -> dict[str, Any] | None:
    """
    Hugging Face NER 異쒕젰(aggregation_strategy="simple" 湲곗?)??Detection dict濡?蹂?섑빀?덈떎.

    PERSON 怨꾩뿴 ?쇰꺼留?蹂?섑빀?덈떎.
    confidence < threshold?대㈃ None??諛섑솚?⑸땲??
    """
    entity_label = (raw.get("entity_group") or raw.get("entity") or "").upper()
    entity_label = entity_label.replace("B-", "").replace("I-", "")

    if entity_label not in {"PERSON", "PER", "PS", "?몃챸"}:
        return None

    score = float(raw.get("score") or 0.0)
    if score < threshold:
        return None

    start = raw.get("start")
    end = raw.get("end")

    if start is None or end is None:
        return None

    start = int(start)
    end = int(end)

    matched = paragraph.text[start:end] or raw.get("word") or ""

    return {
        "label": "?깅챸",
        "matched": str(matched),
        "grade": "S",
        "action": "留덉뒪??,
        "source": "ner",
        "context": paragraph.text,
        "locationLabel": paragraph.location_label,
        "locationMeta": paragraph.location_meta,
        "start": start,
        "end": end,
        "sensitiveType": "媛쒖씤?뺣낫",
        "sensitiveCategory": "?깅챸",
        "reason": (
            f"NER 紐⑤뜽 PERSON ?먯? / original_label={raw.get('entity_group') or raw.get('entity')}"
            f" / confidence={score:.4f} / threshold={threshold:.2f}"
        ),
        "_order": order,
    }


def _make_target_dict_from_ai(
    grade: str,
    confidence: float,
    paragraph: ParsedParagraph,
    order: int,
    *,
    threshold: float,
    prob_map: dict[str, float] | None = None,
) -> dict[str, Any] | None:
    """
    AI 臾몄옣遺꾨쪟 寃곌낵瑜?review target dict濡?蹂?섑빀?덈떎.

    AI Detection? start/end媛 ?녾퀬 matched媛 鍮?臾몄옄?댁엯?덈떎.
    grade=='O'?닿굅??confidence < threshold?대㈃ review target??留뚮뱾吏 ?딆뒿?덈떎.
    """
    if grade == "O" or confidence < threshold:
        return None

    prob_text = ""
    if prob_map:
        prob_text = " / probs=(" + ", ".join(
            f"{label}={prob:.3f}" for label, prob in prob_map.items()
        ) + ")"

    return {
        "label": "誘쇨컧?뺣낫",
        "matched": "",
        "grade": grade,
        "action": "寃???꾩슂",
        "source": "ai",
        "context": paragraph.text,
        "locationLabel": paragraph.location_label,
        "locationMeta": paragraph.location_meta,
        "start": None,
        "end": None,
        "sensitiveType": "臾몃㎘ 湲곕컲 誘쇨컧?뺣낫",
        "sensitiveCategory": f"AI_{grade}",
        "reason": (
            f"AI 臾몄옣遺꾨쪟 grade={grade} / confidence={confidence:.4f}"
            f" / threshold={threshold:.2f}{prob_text}"
        ),
        "_order": order,
    }


# ?? ?먯? ?뚯씠?꾨씪???????????????????????????????????????????????

def detect_in_docx(
    input_path: str,
    *,
    regex_detect_func: Callable[[str], list[Any]] | None = None,
    ner_detect_func: Callable[[str], list[dict[str, Any]]] | None = None,
    ai_predict_func: Callable[[str], tuple[str, float, dict[str, float]]] | None = None,
    ner_threshold: float = 0.8,
    ai_threshold: float = 0.6,
) -> DeidentifyPlan:
    """
    docx ?뚯씪?먯꽌 蹂몃Ц paragraph瑜??쒗쉶?섎ŉ ?먯?瑜??섑뻾?섍퀬 DeidentifyPlan???앹꽦?⑸땲??

    Args:
        input_path: docx ?뚯씪 寃쎈줈
        regex_detect_func: text -> regex detection 紐⑸줉 (?앸왂 ??regex_detector.detect_patterns ?ъ슜)
        ner_detect_func: text -> HF NER pipeline 異쒕젰 紐⑸줉 (?앸왂 ??NER skip)
        ai_predict_func: text -> (grade, confidence, prob_map) (?앸왂 ??AI skip)
        ner_threshold: NER confidence ?꾧퀎媛?        ai_threshold: AI confidence ?꾧퀎媛?
    Returns:
        DeidentifyPlan (auto_targets + review_targets)

    ?먯? ?⑥닔瑜?二쇱엯?뺤쑝濡?諛쏅뒗 ?댁쑀:
    - ?⑥쐞 ?뚯뒪?몄뿉??紐⑤뜽 ?섏〈?깆쓣 ?딆쓣 ???덈룄濡??⑸땲??
    - 13二쇱감 珥덈컲?먮뒗 regex留??곌껐?댁꽌 guide 援ъ“遺??寃利앺븷 ???덉뒿?덈떎.
    """
    # regex ?먯? ?⑥닔 湲곕낯媛?    if regex_detect_func is None:
        try:
            from src.regex_detector import detect_patterns as _detect_patterns
        except ModuleNotFoundError:
            from regex_detector import detect_patterns as _detect_patterns
        regex_detect_func = _detect_patterns

    doc = load_docx(input_path)
    paragraphs = iter_docx_paragraphs(doc)

    detections: list[dict[str, Any]] = []
    order = 0

    for paragraph in paragraphs:
        # regex
        raw_regex = regex_detect_func(paragraph.text) or []
        for raw in raw_regex:
            detection = _make_target_dict_from_regex(raw, paragraph, order)
            if detection is not None:
                detections.append(detection)
                order += 1

        # NER
        if ner_detect_func is not None:
            try:
                raw_ner = ner_detect_func(paragraph.text) or []
            except Exception as exc:
                print(f"[NER] {paragraph.location_label} ?먯? ?ㅽ뙣: {exc}")
                raw_ner = []

            for raw in raw_ner:
                detection = _make_target_dict_from_ner(
                    raw, paragraph, order, threshold=ner_threshold,
                )
                if detection is not None:
                    detections.append(detection)
                    order += 1

        # AI
        if ai_predict_func is not None:
            try:
                grade, confidence, prob_map = ai_predict_func(paragraph.text)
            except Exception as exc:
                print(f"[AI] {paragraph.location_label} ?덉륫 ?ㅽ뙣: {exc}")
                grade, confidence, prob_map = "O", 0.0, {}

            if grade is not None and confidence is not None:
                detection = _make_target_dict_from_ai(
                    grade, confidence, paragraph, order,
                    threshold=ai_threshold, prob_map=prob_map,
                )
                if detection is not None:
                    detections.append(detection)
                    order += 1

    return build_deidentify_plan(detections)


# ?? guide ?앹꽦 ?????????????????????????????????????????????????

def _make_skipped_item_for_target(
    target: DeidentifyTarget,
    warning_type: str,
    message: str,
) -> CommonApplyItem:
    warning = format_warning(warning_type, message)
    return CommonApplyItem(
        locationLabel=target.location_label,
        locationMeta=target.location_meta or {},
        label=target.label or "",
        action=target.action,
        originalText=target.context or "",
        appliedText=target.context or "",
        status="skipped",
        appliedTargetCount=0,
        skippedTargetCount=1,
        warnings=[warning],
    )


def _target_location_key(target: DeidentifyTarget) -> tuple | None:
    meta = target.location_meta or {}
    section = str(meta.get("section") or "body")
    paragraph_no = meta.get("paragraphNo")

    if paragraph_no is None:
        return None

    if section == "body":
        return ("body", int(paragraph_no))

    if section == "table_cell":
        table_no = meta.get("tableNo")
        row_no = meta.get("rowNo")
        col_no = meta.get("colNo")
        if table_no is None or row_no is None or col_no is None:
            return None
        return ("table_cell", int(table_no), int(row_no), int(col_no), int(paragraph_no))

    return None


def _group_targets_by_location(
    targets: list[DeidentifyTarget],
) -> tuple[dict[tuple, list[DeidentifyTarget]], list[CommonApplyItem], list[str]]:
    """
    auto target??docx location key 湲곗??쇰줈 臾띠뒿?덈떎.

    body key:
        ("body", paragraphNo)
    table_cell key:
        ("table_cell", tableNo, rowNo, colNo, paragraphNo)
    """
    grouped: dict[tuple, list[DeidentifyTarget]] = {}
    skipped_items: list[CommonApplyItem] = []
    warnings: list[str] = []

    for target in targets:
        meta = target.location_meta or {}

        if str(meta.get("fileType") or "").lower() != "docx":
            continue

        section = str(meta.get("section") or "body")
        paragraph_no = meta.get("paragraphNo")

        if paragraph_no is None:
            item = _make_skipped_item_for_target(
                target,
                WARNING_MISSING_PARAGRAPH_NO,
                f"{target.location_label}: paragraphNo媛 ?놁뼱 ?덈궡瑜??앹꽦?섏? 紐삵뻽?듬땲??",
            )
            warnings.extend(item.warnings)
            skipped_items.append(item)
            continue

        if section not in {"body", "table_cell"}:
            item = _make_skipped_item_for_target(
                target,
                WARNING_UNSUPPORTED_DOCX_SECTION,
                f"{target.location_label}: section={section} ?꾩튂???꾩옱 docx guide 踰붿쐞 ?몄씠誘濡??덈궡瑜??앹꽦?섏? ?딆뒿?덈떎.",
            )
            warnings.extend(item.warnings)
            skipped_items.append(item)
            continue

        key = _target_location_key(target)
        if key is None:
            item = _make_skipped_item_for_target(
                target,
                WARNING_MISSING_PARAGRAPH_NO,
                f"{target.location_label}: ???꾩튂 硫뷀??곗씠??tableNo/rowNo/colNo/paragraphNo)媛 遺議깊빐 ?덈궡瑜??앹꽦?섏? 紐삵뻽?듬땲??",
            )
            warnings.extend(item.warnings)
            skipped_items.append(item)
            continue

        grouped.setdefault(key, []).append(target)

    return grouped, skipped_items, warnings


def _build_guide_item_for_paragraph(
    parsed_paragraph: ParsedParagraph | None,
    location_key: tuple,
    targets: list[DeidentifyTarget],
    *,
    deletion_mode: str,
) -> CommonApplyItem:
    """
    ??paragraph???랁븳 target 紐⑸줉?????guide 紐⑤뱶 CommonApplyItem???앹꽦?⑸땲??
    """
    representative = targets[0]

    if parsed_paragraph is not None:
        paragraph_text = parsed_paragraph.text
        location_label = representative.location_label or parsed_paragraph.location_label
        location_meta = representative.location_meta or parsed_paragraph.location_meta
    else:
        paragraph_text = None
        location_label = representative.location_label or "?????녿뒗 docx ?꾩튂"
        location_meta = representative.location_meta or {"fileType": "docx"}

    warnings: list[str] = []

    # paragraph_text媛 None?대㈃ paragraph 踰붿쐞 珥덇낵
    if paragraph_text is None:
        warning = format_warning(
            WARNING_PARAGRAPH_OUT_OF_RANGE,
            f"{location_label}: location={location_key}媛 臾몄꽌 踰붿쐞瑜?踰쀬뼱?ъ뒿?덈떎.",
        )
        warnings.append(warning)
        return CommonApplyItem(
            locationLabel=location_label,
            locationMeta=location_meta,
            label=labels_for_targets(targets),
            action=actions_for_targets(targets),
            originalText="",
            appliedText="",
            status="skipped",
            appliedTargetCount=0,
            skippedTargetCount=len(targets),
            warnings=warnings,
        )

    # 鍮?paragraph
    if not paragraph_text.strip():
        warning = format_warning(
            WARNING_EMPTY_PARAGRAPH_TARGET,
            f"{location_label}: 鍮?paragraph瑜?媛由ы궎??target? ?덈궡瑜??앹꽦?섏? ?딆뒿?덈떎.",
        )
        warnings.append(warning)
        return CommonApplyItem(
            locationLabel=location_label,
            locationMeta=location_meta,
            label=labels_for_targets(targets),
            action=actions_for_targets(targets),
            originalText="",
            appliedText="",
            status="skipped",
            appliedTargetCount=0,
            skippedTargetCount=len(targets),
            warnings=warnings,
        )

    # context 遺덉씪移?(?곸슜? 吏꾪뻾)
    if any(
        target.context is not None
        and normalize_nfc(target.context) != normalize_nfc(paragraph_text)
        for target in targets
    ):
        warnings.append(
            format_warning(
                WARNING_CONTEXT_MISMATCH,
                f"{location_label}: target.context? ?ㅼ젣 paragraph ?띿뒪?멸? ?ㅻ쫭?덈떎. "
                "paragraph ?띿뒪??湲곗??쇰줈 slice 寃利???沅뚯옣 ?щ?瑜??먮떒?⑸땲??",
            )
        )

    # slice 寃利?    valid_targets: list[DeidentifyTarget] = []
    skipped_count = 0

    for target in targets:
        warning_type, slice_error = validate_slice_against_text(paragraph_text, target)
        if slice_error is not None:
            warnings.append(
                format_warning(warning_type, f"{location_label}: {slice_error}")
            )
            skipped_count += 1
            continue

        valid_targets.append(target)

    # guide 紐⑤뱶 preview ?앹꽦 (硫붾え由ъ뿉?쒕쭔)
    if valid_targets:
        apply_result = apply_targets_to_text(
            paragraph_text,
            valid_targets,
            deletion_mode=deletion_mode,
        )
        applied_text = apply_result.applied_text
        warnings.extend(apply_result.warnings)
        applied_count = len(apply_result.applied_targets)
        skipped_count += len(apply_result.skipped_targets)
    else:
        applied_text = paragraph_text
        applied_count = 0

    return CommonApplyItem(
        locationLabel=location_label,
        locationMeta=location_meta,
        label=labels_for_targets(targets),
        action=actions_for_targets(targets),
        originalText=paragraph_text,
        appliedText=applied_text,
        status=make_status(applied_count, skipped_count),
        appliedTargetCount=applied_count,
        skippedTargetCount=skipped_count,
        warnings=warnings,
    )


def _paragraph_location_key(paragraph: ParsedParagraph) -> tuple:
    if paragraph.section == "table_cell":
        return (
            "table_cell",
            paragraph.table_no,
            paragraph.row_no,
            paragraph.col_no,
            paragraph.paragraph_no,
        )
    return ("body", paragraph.paragraph_no)


def build_guide_for_docx(
    input_path: str,
    plan: DeidentifyPlan,
    *,
    deletion_mode: str = "delete",
) -> CommonApplyResult:
    """
    DeidentifyPlan??諛쏆븘 guide 紐⑤뱶 CommonApplyResult瑜??앹꽦?⑸땲??

    ?ㅼ젣 ?뚯씪???섏젙?섏? ?딆쑝硫? outputFilePath??None?낅땲??
    """
    doc = load_docx(input_path)
    parsed_paragraphs = iter_docx_paragraphs(doc)
    paragraph_map = {
        _paragraph_location_key(parsed): parsed
        for parsed in parsed_paragraphs
    }

    grouped, skipped_items, global_warnings = _group_targets_by_location(plan.auto_targets)

    auto_results: list[CommonApplyItem] = list(skipped_items)

    for location_key, targets in grouped.items():
        parsed = paragraph_map.get(location_key)

        item = _build_guide_item_for_paragraph(
            parsed,
            location_key,
            targets,
            deletion_mode=deletion_mode,
        )
        auto_results.append(item)

    review_items = make_review_items(plan.review_targets)
    summary = build_summary(auto_results, review_items, global_warnings)

    return CommonApplyResult(
        fileType="docx",
        applyMode=APPLY_MODE_GUIDE,
        inputFilePath=str(input_path),
        outputFilePath=None,
        autoResults=auto_results,
        reviewTargets=review_items,
        warnings=global_warnings,
        summary=summary,
    )


def detect_and_build_guide_for_docx(
    input_path: str,
    *,
    regex_detect_func: Callable[[str], list[Any]] | None = None,
    ner_detect_func: Callable[[str], list[dict[str, Any]]] | None = None,
    ai_predict_func: Callable[[str], tuple[str, float, dict[str, float]]] | None = None,
    ner_threshold: float = 0.8,
    ai_threshold: float = 0.6,
    deletion_mode: str = "delete",
) -> CommonApplyResult:
    """
    detect_in_docx + build_guide_for_docx ?몄쓽 wrapper.
    """
    plan = detect_in_docx(
        input_path,
        regex_detect_func=regex_detect_func,
        ner_detect_func=ner_detect_func,
        ai_predict_func=ai_predict_func,
        ner_threshold=ner_threshold,
        ai_threshold=ai_threshold,
    )

    return build_guide_for_docx(
        input_path,
        plan,
        deletion_mode=deletion_mode,
    )
