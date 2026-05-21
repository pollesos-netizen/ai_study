"""
docx 파일 탐지 + 안내(guide) 모드

목적:
- docx 파일에서 본문 paragraph를 순회하며 개인정보/민감정보를 탐지합니다.
- 시스템은 docx 파일을 직접 수정하지 않습니다.
- 결과는 CommonApplyResult(applyMode="guide")로 반환합니다.
- 사용자는 안내에 따라 원본 docx에서 직접 수정합니다.

13주차 범위:
- 본문 paragraphs만 처리 (표/헤더/푸터/각주는 별도 주차)
- 빈 문단(strip 기준)은 탐지 대상에서 제외
- regex + NER + AI 탐지를 지원하되, 13주차 초반에는 regex만 우선 검증 가능

핵심 함수:
- detect_in_docx():            DeidentifyPlan 생성
- build_guide_for_docx():      DeidentifyPlan -> CommonApplyResult (guide)
- detect_and_build_guide_for_docx(): 두 함수의 편의 wrapper
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

try:
    from src.common_apply_result import (
        APPLY_MODE_GUIDE,
        CommonApplyItem,
        CommonApplyResult,
        build_summary,
        make_review_items,
    )
    from src.common_apply_utils import (
        WARNING_CONTEXT_MISMATCH,
        WARNING_EMPTY_PARAGRAPH_TARGET,
        WARNING_MISSING_PARAGRAPH_NO,
        WARNING_PARAGRAPH_NOT_IN_BODY,
        WARNING_PARAGRAPH_OUT_OF_RANGE,
        actions_for_targets,
        format_warning,
        labels_for_targets,
        make_location_label_with_context,
        make_status,
        normalize_nfc,
        validate_slice_against_text,
    )
    from src.deidentify_apply import apply_targets_to_text
    from src.deidentify_target_builder import (
        DeidentifyPlan,
        DeidentifyTarget,
        build_deidentify_plan,
    )
except ModuleNotFoundError:
    from common_apply_result import (
        APPLY_MODE_GUIDE,
        CommonApplyItem,
        CommonApplyResult,
        build_summary,
        make_review_items,
    )
    from common_apply_utils import (
        WARNING_CONTEXT_MISMATCH,
        WARNING_EMPTY_PARAGRAPH_TARGET,
        WARNING_MISSING_PARAGRAPH_NO,
        WARNING_PARAGRAPH_NOT_IN_BODY,
        WARNING_PARAGRAPH_OUT_OF_RANGE,
        actions_for_targets,
        format_warning,
        labels_for_targets,
        make_location_label_with_context,
        make_status,
        normalize_nfc,
        validate_slice_against_text,
    )
    from deidentify_apply import apply_targets_to_text
    from deidentify_target_builder import (
        DeidentifyPlan,
        DeidentifyTarget,
        build_deidentify_plan,
    )


# ── 데이터 구조 ────────────────────────────────────────────────

@dataclass
class ParsedParagraph:
    """
    docx의 paragraph를 우리 탐지 단위로 변환한 구조.

    section:
    - "body": doc.paragraphs 기준 본문 paragraph
    - "table_cell": doc.tables 내부 cell.paragraphs 기준 paragraph
    """

    paragraph_no: int  # section별 paragraph 인덱스 (0-based)
    section: str
    text: str
    table_no: int | None = None
    row_no: int | None = None
    col_no: int | None = None

    @property
    def location_label(self) -> str:
        # 사용자 표시만 1-based로 변환 + context 30자
        if self.section == "table_cell":
            base = (
                f"표 {self.table_no + 1}번 "
                f"{self.row_no + 1}행 {self.col_no + 1}열"
            )
            if self.paragraph_no > 0:
                base += f" {self.paragraph_no + 1}번째 문단"
        else:
            base = f"본문 {self.paragraph_no + 1}번째 문단"

        return make_location_label_with_context(base, self.text, max_length=30)

    @property
    def location_meta(self) -> dict[str, Any]:
        meta = {
            "fileType": "docx",
            "section": self.section,
            "paragraphNo": self.paragraph_no,
        }

        if self.section == "table_cell":
            meta.update({
                "tableNo": self.table_no,
                "rowNo": self.row_no,
                "colNo": self.col_no,
            })

        return meta


# ── docx 로드 및 paragraph 순회 ────────────────────────────────

def load_docx(input_path: str | Path):
    """
    python-docx로 docx 파일을 로드합니다.

    의존성: python-docx (`pip install python-docx`)
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx가 설치되어 있지 않습니다. "
            "`pip install python-docx`를 실행하세요."
        ) from exc

    return Document(str(input_path))


def iter_body_paragraphs(doc) -> list[ParsedParagraph]:
    """
    문서 본문(doc.paragraphs)의 paragraph를 ParsedParagraph 목록으로 반환합니다.

    - 빈 문단(strip() 기준)은 제외합니다.
    - paragraphNo는 빈 문단을 포함한 원문 인덱스를 유지합니다.
    - 13주차에서는 본문만 처리하므로 section="body"로 고정합니다.
    """
    parsed: list[ParsedParagraph] = []

    for para_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text

        if not text.strip():
            continue

        parsed.append(
            ParsedParagraph(
                paragraph_no=para_index,
                section="body",
                text=text,
            )
        )

    return parsed

def iter_table_cell_paragraphs(doc) -> list[ParsedParagraph]:
    """
    문서 표 셀 내부 paragraph를 ParsedParagraph 목록으로 반환합니다.

    - 빈 paragraph(strip 기준)는 제외합니다.
    - tableNo/rowNo/colNo/paragraphNo는 0-based로 저장합니다.
    - 병합 셀 중복 제거는 13주차에서는 수행하지 않습니다.
      guide 모드에서는 누락 방지가 중복 제거보다 중요합니다.
    """
    parsed: list[ParsedParagraph] = []

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text

                    if not text.strip():
                        continue

                    parsed.append(
                        ParsedParagraph(
                            paragraph_no=para_index,
                            section="table_cell",
                            text=text,
                            table_no=table_index,
                            row_no=row_index,
                            col_no=col_index,
                        )
                    )

    return parsed

def iter_docx_paragraphs(doc) -> list[ParsedParagraph]:
    """
    13주차 docx guide 탐지 대상 paragraph를 반환합니다.

    대상:
    - 본문 paragraph
    - 표 셀 내부 paragraph

    제외:
    - 헤더/푸터/각주/주석/도형/SmartArt/차트 내부 텍스트
    """
    return iter_body_paragraphs(doc) + iter_table_cell_paragraphs(doc)


# ── Detection 생성 (regex / NER / AI 어댑터) ───────────────────

def _make_target_dict_from_regex(
    raw: Any,
    paragraph: ParsedParagraph,
    order: int,
) -> dict[str, Any] | None:
    """
    regex_detector의 결과(DetectionResult 또는 유사 dict)를 Detection dict로 변환합니다.

    DeidentifyPlan 생성기는 dict 입력을 받으므로 dict 형태로 정규화합니다.
    """
    def _get(obj, *names, default=None):
        for name in names:
            if isinstance(obj, dict) and name in obj:
                return obj.get(name)
            if hasattr(obj, name):
                return getattr(obj, name)
        return default

    label = _get(raw, "label")
    value = _get(raw, "value", "matched", "match")
    start = _get(raw, "start")
    end = _get(raw, "end")
    grade = _get(raw, "grade", default="S")
    action = _get(raw, "action", default="마스킹")
    desc = _get(raw, "desc", "reason", default=None)

    if label is None or value is None or start is None or end is None:
        return None

    return {
        "label": str(label),
        "matched": str(value),
        "grade": str(grade),
        "action": str(action),
        "source": "regex",
        "context": paragraph.text,
        "locationLabel": paragraph.location_label,
        "locationMeta": paragraph.location_meta,
        "start": int(start),
        "end": int(end),
        "sensitiveType": _get(raw, "sensitive_type", "sensitiveType", default=None),
        "sensitiveCategory": _get(raw, "sensitive_category", "sensitiveCategory", default=label),
        "reason": str(desc) if desc else f"정규식 탐지: {label}",
        "_order": order,
    }


def _make_target_dict_from_ner(
    raw: dict[str, Any],
    paragraph: ParsedParagraph,
    order: int,
    *,
    threshold: float,
) -> dict[str, Any] | None:
    """
    Hugging Face NER 출력(aggregation_strategy="simple" 기준)을 Detection dict로 변환합니다.

    PERSON 계열 라벨만 변환합니다.
    confidence < threshold이면 None을 반환합니다.
    """
    entity_label = (raw.get("entity_group") or raw.get("entity") or "").upper()
    entity_label = entity_label.replace("B-", "").replace("I-", "")

    if entity_label not in {"PERSON", "PER", "PS", "인명"}:
        return None

    score = float(raw.get("score") or 0.0)
    if score < threshold:
        return None

    start = raw.get("start")
    end = raw.get("end")

    if start is None or end is None:
        return None

    start = int(start)
    end = int(end)

    matched = paragraph.text[start:end] or raw.get("word") or ""

    return {
        "label": "성명",
        "matched": str(matched),
        "grade": "S",
        "action": "마스킹",
        "source": "ner",
        "context": paragraph.text,
        "locationLabel": paragraph.location_label,
        "locationMeta": paragraph.location_meta,
        "start": start,
        "end": end,
        "sensitiveType": "개인정보",
        "sensitiveCategory": "성명",
        "reason": (
            f"NER 모델 PERSON 탐지 / original_label={raw.get('entity_group') or raw.get('entity')}"
            f" / confidence={score:.4f} / threshold={threshold:.2f}"
        ),
        "_order": order,
    }


def _make_target_dict_from_ai(
    grade: str,
    confidence: float,
    paragraph: ParsedParagraph,
    order: int,
    *,
    threshold: float,
    prob_map: dict[str, float] | None = None,
) -> dict[str, Any] | None:
    """
    AI 문장분류 결과를 review target dict로 변환합니다.

    AI Detection은 start/end가 없고 matched가 빈 문자열입니다.
    grade=='O'이거나 confidence < threshold이면 review target을 만들지 않습니다.
    """
    if grade == "O" or confidence < threshold:
        return None

    prob_text = ""
    if prob_map:
        prob_text = " / probs=(" + ", ".join(
            f"{label}={prob:.3f}" for label, prob in prob_map.items()
        ) + ")"

    return {
        "label": "민감정보",
        "matched": "",
        "grade": grade,
        "action": "검토 필요",
        "source": "ai",
        "context": paragraph.text,
        "locationLabel": paragraph.location_label,
        "locationMeta": paragraph.location_meta,
        "start": None,
        "end": None,
        "sensitiveType": "문맥 기반 민감정보",
        "sensitiveCategory": f"AI_{grade}",
        "reason": (
            f"AI 문장분류 grade={grade} / confidence={confidence:.4f}"
            f" / threshold={threshold:.2f}{prob_text}"
        ),
        "_order": order,
    }


# ── 탐지 파이프라인 ─────────────────────────────────────────────

def detect_in_docx(
    input_path: str,
    *,
    regex_detect_func: Callable[[str], list[Any]] | None = None,
    ner_detect_func: Callable[[str], list[dict[str, Any]]] | None = None,
    ai_predict_func: Callable[[str], tuple[str, float, dict[str, float]]] | None = None,
    ner_threshold: float = 0.8,
    ai_threshold: float = 0.6,
) -> DeidentifyPlan:
    """
    docx 파일에서 본문 paragraph를 순회하며 탐지를 수행하고 DeidentifyPlan을 생성합니다.

    Args:
        input_path: docx 파일 경로
        regex_detect_func: text -> regex detection 목록 (생략 시 regex_detector.detect_patterns 사용)
        ner_detect_func: text -> HF NER pipeline 출력 목록 (생략 시 NER skip)
        ai_predict_func: text -> (grade, confidence, prob_map) (생략 시 AI skip)
        ner_threshold: NER confidence 임계값
        ai_threshold: AI confidence 임계값

    Returns:
        DeidentifyPlan (auto_targets + review_targets)

    탐지 함수를 주입형으로 받는 이유:
    - 단위 테스트에서 모델 의존성을 끊을 수 있도록 합니다.
    - 13주차 초반에는 regex만 연결해서 guide 구조부터 검증할 수 있습니다.
    """
    # regex 탐지 함수 기본값
    if regex_detect_func is None:
        try:
            from src.regex_detector import detect_patterns as _detect_patterns
        except ModuleNotFoundError:
            from regex_detector import detect_patterns as _detect_patterns
        regex_detect_func = _detect_patterns

    doc = load_docx(input_path)
    paragraphs = iter_docx_paragraphs(doc)

    detections: list[dict[str, Any]] = []
    order = 0

    for paragraph in paragraphs:
        # regex
        raw_regex = regex_detect_func(paragraph.text) or []
        for raw in raw_regex:
            detection = _make_target_dict_from_regex(raw, paragraph, order)
            if detection is not None:
                detections.append(detection)
                order += 1

        # NER
        if ner_detect_func is not None:
            try:
                raw_ner = ner_detect_func(paragraph.text) or []
            except Exception as exc:
                print(f"[NER] {paragraph.location_label} 탐지 실패: {exc}")
                raw_ner = []

            for raw in raw_ner:
                detection = _make_target_dict_from_ner(
                    raw, paragraph, order, threshold=ner_threshold,
                )
                if detection is not None:
                    detections.append(detection)
                    order += 1

        # AI
        if ai_predict_func is not None:
            try:
                grade, confidence, prob_map = ai_predict_func(paragraph.text)
            except Exception as exc:
                print(f"[AI] {paragraph.location_label} 예측 실패: {exc}")
                grade, confidence, prob_map = "O", 0.0, {}

            if grade is not None and confidence is not None:
                detection = _make_target_dict_from_ai(
                    grade, confidence, paragraph, order,
                    threshold=ai_threshold, prob_map=prob_map,
                )
                if detection is not None:
                    detections.append(detection)
                    order += 1

    return build_deidentify_plan(detections)


# ── guide 생성 ─────────────────────────────────────────────────

def _make_skipped_item_for_target(
    target: DeidentifyTarget,
    warning_type: str,
    message: str,
) -> CommonApplyItem:
    warning = format_warning(warning_type, message)
    return CommonApplyItem(
        locationLabel=target.location_label,
        locationMeta=target.location_meta or {},
        label=target.label or "",
        action=target.action,
        originalText=target.context or "",
        appliedText=target.context or "",
        status="skipped",
        appliedTargetCount=0,
        skippedTargetCount=1,
        warnings=[warning],
    )


def _target_location_key(target: DeidentifyTarget) -> tuple | None:
    meta = target.location_meta or {}
    section = str(meta.get("section") or "body")
    paragraph_no = meta.get("paragraphNo")

    if paragraph_no is None:
        return None

    if section == "body":
        return ("body", int(paragraph_no))

    if section == "table_cell":
        table_no = meta.get("tableNo")
        row_no = meta.get("rowNo")
        col_no = meta.get("colNo")
        if table_no is None or row_no is None or col_no is None:
            return None
        return ("table_cell", int(table_no), int(row_no), int(col_no), int(paragraph_no))

    return None


def _group_targets_by_location(
    targets: list[DeidentifyTarget],
) -> tuple[dict[tuple, list[DeidentifyTarget]], list[CommonApplyItem], list[str]]:
    """
    auto target을 docx location key 기준으로 묶습니다.

    body key:
        ("body", paragraphNo)
    table_cell key:
        ("table_cell", tableNo, rowNo, colNo, paragraphNo)
    """
    grouped: dict[tuple, list[DeidentifyTarget]] = {}
    skipped_items: list[CommonApplyItem] = []
    warnings: list[str] = []

    for target in targets:
        meta = target.location_meta or {}

        if str(meta.get("fileType") or "").lower() != "docx":
            continue

        section = str(meta.get("section") or "body")
        paragraph_no = meta.get("paragraphNo")

        if paragraph_no is None:
            item = _make_skipped_item_for_target(
                target,
                WARNING_MISSING_PARAGRAPH_NO,
                f"{target.location_label}: paragraphNo가 없어 안내를 생성하지 못했습니다.",
            )
            warnings.extend(item.warnings)
            skipped_items.append(item)
            continue

        if section not in {"body", "table_cell"}:
            item = _make_skipped_item_for_target(
                target,
                WARNING_PARAGRAPH_NOT_IN_BODY,
                f"{target.location_label}: section={section} 위치는 현재 docx guide 범위 외이므로 안내를 생성하지 않습니다.",
            )
            warnings.extend(item.warnings)
            skipped_items.append(item)
            continue

        key = _target_location_key(target)
        if key is None:
            item = _make_skipped_item_for_target(
                target,
                WARNING_MISSING_PARAGRAPH_NO,
                f"{target.location_label}: 표 위치 메타데이터(tableNo/rowNo/colNo/paragraphNo)가 부족해 안내를 생성하지 못했습니다.",
            )
            warnings.extend(item.warnings)
            skipped_items.append(item)
            continue

        grouped.setdefault(key, []).append(target)

    return grouped, skipped_items, warnings


def _build_guide_item_for_paragraph(
    parsed_paragraph: ParsedParagraph | None,
    location_key: tuple,
    targets: list[DeidentifyTarget],
    *,
    deletion_mode: str,
) -> CommonApplyItem:
    """
    한 paragraph에 속한 target 목록에 대해 guide 모드 CommonApplyItem을 생성합니다.
    """
    representative = targets[0]

    if parsed_paragraph is not None:
        paragraph_text = parsed_paragraph.text
        location_label = representative.location_label or parsed_paragraph.location_label
        location_meta = representative.location_meta or parsed_paragraph.location_meta
    else:
        paragraph_text = None
        location_label = representative.location_label or "알 수 없는 docx 위치"
        location_meta = representative.location_meta or {"fileType": "docx"}

    warnings: list[str] = []

    # paragraph_text가 None이면 paragraph 범위 초과
    if paragraph_text is None:
        warning = format_warning(
            WARNING_PARAGRAPH_OUT_OF_RANGE,
            f"{location_label}: location={location_key}가 문서 범위를 벗어났습니다.",
        )
        warnings.append(warning)
        return CommonApplyItem(
            locationLabel=location_label,
            locationMeta=location_meta,
            label=labels_for_targets(targets),
            action=actions_for_targets(targets),
            originalText="",
            appliedText="",
            status="skipped",
            appliedTargetCount=0,
            skippedTargetCount=len(targets),
            warnings=warnings,
        )

    # 빈 paragraph
    if not paragraph_text.strip():
        warning = format_warning(
            WARNING_EMPTY_PARAGRAPH_TARGET,
            f"{location_label}: 빈 paragraph를 가리키는 target은 안내를 생성하지 않습니다.",
        )
        warnings.append(warning)
        return CommonApplyItem(
            locationLabel=location_label,
            locationMeta=location_meta,
            label=labels_for_targets(targets),
            action=actions_for_targets(targets),
            originalText="",
            appliedText="",
            status="skipped",
            appliedTargetCount=0,
            skippedTargetCount=len(targets),
            warnings=warnings,
        )

    # context 불일치 (적용은 진행)
    if any(
        target.context is not None
        and normalize_nfc(target.context) != normalize_nfc(paragraph_text)
        for target in targets
    ):
        warnings.append(
            format_warning(
                WARNING_CONTEXT_MISMATCH,
                f"{location_label}: target.context와 실제 paragraph 텍스트가 다릅니다. "
                "paragraph 텍스트 기준으로 slice 검증 후 권장 여부를 판단합니다.",
            )
        )

    # slice 검증
    valid_targets: list[DeidentifyTarget] = []
    skipped_count = 0

    for target in targets:
        warning_type, slice_error = validate_slice_against_text(paragraph_text, target)
        if slice_error is not None:
            warnings.append(
                format_warning(warning_type, f"{location_label}: {slice_error}")
            )
            skipped_count += 1
            continue

        valid_targets.append(target)

    # guide 모드 preview 생성 (메모리에서만)
    if valid_targets:
        apply_result = apply_targets_to_text(
            paragraph_text,
            valid_targets,
            deletion_mode=deletion_mode,
        )
        applied_text = apply_result.applied_text
        warnings.extend(apply_result.warnings)
        applied_count = len(apply_result.applied_targets)
        skipped_count += len(apply_result.skipped_targets)
    else:
        applied_text = paragraph_text
        applied_count = 0

    return CommonApplyItem(
        locationLabel=location_label,
        locationMeta=location_meta,
        label=labels_for_targets(targets),
        action=actions_for_targets(targets),
        originalText=paragraph_text,
        appliedText=applied_text,
        status=make_status(applied_count, skipped_count),
        appliedTargetCount=applied_count,
        skippedTargetCount=skipped_count,
        warnings=warnings,
    )


def _paragraph_location_key(paragraph: ParsedParagraph) -> tuple:
    if paragraph.section == "table_cell":
        return (
            "table_cell",
            paragraph.table_no,
            paragraph.row_no,
            paragraph.col_no,
            paragraph.paragraph_no,
        )
    return ("body", paragraph.paragraph_no)


def build_guide_for_docx(
    input_path: str,
    plan: DeidentifyPlan,
    *,
    deletion_mode: str = "delete",
) -> CommonApplyResult:
    """
    DeidentifyPlan을 받아 guide 모드 CommonApplyResult를 생성합니다.

    실제 파일을 수정하지 않으며, outputFilePath는 None입니다.
    """
    doc = load_docx(input_path)
    parsed_paragraphs = iter_docx_paragraphs(doc)
    paragraph_map = {
        _paragraph_location_key(parsed): parsed
        for parsed in parsed_paragraphs
    }

    grouped, skipped_items, global_warnings = _group_targets_by_location(plan.auto_targets)

    auto_results: list[CommonApplyItem] = list(skipped_items)

    for location_key, targets in grouped.items():
        parsed = paragraph_map.get(location_key)

        item = _build_guide_item_for_paragraph(
            parsed,
            location_key,
            targets,
            deletion_mode=deletion_mode,
        )
        auto_results.append(item)

    review_items = make_review_items(plan.review_targets)
    summary = build_summary(auto_results, review_items, global_warnings)

    return CommonApplyResult(
        fileType="docx",
        applyMode=APPLY_MODE_GUIDE,
        inputFilePath=str(input_path),
        outputFilePath=None,
        autoResults=auto_results,
        reviewTargets=review_items,
        warnings=global_warnings,
        summary=summary,
    )


def detect_and_build_guide_for_docx(
    input_path: str,
    *,
    regex_detect_func: Callable[[str], list[Any]] | None = None,
    ner_detect_func: Callable[[str], list[dict[str, Any]]] | None = None,
    ai_predict_func: Callable[[str], tuple[str, float, dict[str, float]]] | None = None,
    ner_threshold: float = 0.8,
    ai_threshold: float = 0.6,
    deletion_mode: str = "delete",
) -> CommonApplyResult:
    """
    detect_in_docx + build_guide_for_docx 편의 wrapper.
    """
    plan = detect_in_docx(
        input_path,
        regex_detect_func=regex_detect_func,
        ner_detect_func=ner_detect_func,
        ai_predict_func=ai_predict_func,
        ner_threshold=ner_threshold,
        ai_threshold=ai_threshold,
    )

    return build_guide_for_docx(
        input_path,
        plan,
        deletion_mode=deletion_mode,
    )
